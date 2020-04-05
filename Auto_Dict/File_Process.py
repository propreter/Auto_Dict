# coding=utf-8
import os
import re

import docx
from docx.enum.text import WD_COLOR_INDEX

# 判断文本字符
regexp_char = r'[A-Za-z]+'


# 拷贝源文档内容
def get_para_data(input_doc, paragraph):
    """
    Write the run to the new file and then set its font, bold, alignment, color etc. data.
    """
    output_para = input_doc.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's highlight data
        output_run.font.highlight_color = run.font.highlight_color
        # Run's font data
        output_run.style.name = run.style.name
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment


# 文本预处理：复制源文档到缓存文件
def pre_process(src_path, cache_path):
    cache_name = 'cache'
    cache_ext = '.docx'
    # 打开源文件
    src_docx = docx.Document(src_path)
    # 创建缓存文件
    cache_docx = docx.Document()

    for para in src_docx.paragraphs:
        get_para_data(cache_docx, para)

    # 生成缓存文件
    cache_file = os.path.join(cache_path, cache_name)
    cache_docx.save(cache_file + cache_ext)

    return cache_docx


# 辨别字符类型：中文或英文
def char_discriminator(char):
    pattern = re.compile(regexp_char)
    if pattern.search(char, re.I) is None:
        # 中文
        char_index = True
    else:
        # 英文
        char_index = False
    return char_index


# 从docx文档中提取出黄色高亮文本
def word_extraction(cache_file):
    cnt = 0
    global key_text, chn_text, eng_text
    key_text = []
    chn_text = []
    eng_text = []

    for para in cache_file.paragraphs:
        for run in para.runs:
            if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                key_text.append(run.text.strip())

    for key in key_text:
        if '' in key_text:
            key_text.remove('')

    while cnt < len(key_text):
        for key in key_text:
            # 利用正则识别中英文文本
            if char_discriminator(key):
                chn_text.append(key_text[cnt])
            else:
                eng_text.append(key_text[cnt])
            cnt += 1
    return key_text


# 程序入口
def file_process(src_file, cache_path):
    # 生成缓存文件
    cache = pre_process(src_file, cache_path)
    # 提取指定单词
    word_list = word_extraction(cache)

    return word_list
